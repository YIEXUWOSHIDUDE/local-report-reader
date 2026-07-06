from __future__ import annotations

import re
import shutil
import subprocess
import tempfile
from dataclasses import dataclass
from pathlib import Path

import pdfplumber
import fitz
import pytesseract
from docx import Document
from PIL import Image


@dataclass
class ParsedDocument:
    text: str
    language: str
    notes: list[str]


def parse_document(path: Path, filename: str) -> ParsedDocument:
    suffix = path.suffix.lower()
    notes: list[str] = []
    try:
        if suffix == ".doc":
            text = _parse_legacy_word(path, notes)
        elif suffix == ".docx":
            text = _parse_word(path, notes)
        elif suffix == ".pdf":
            text = _parse_pdf(path, notes)
        elif suffix in {".png", ".jpg", ".jpeg", ".webp", ".tif", ".tiff", ".bmp"}:
            text = _ocr_image(path, notes)
        else:
            notes.append(f"暂不支持的文件类型：{suffix or filename}")
            text = ""
    except Exception as exc:
        notes.append(f"解析失败：{exc}")
        text = ""

    cleaned = _clean_text(text)
    if not cleaned:
        notes.append("未能提取到可用文本，请确认文件是否为加密文件或图片质量是否足够。")
    return ParsedDocument(text=cleaned, language=detect_language(cleaned), notes=notes)


def _parse_legacy_word(path: Path, notes: list[str]) -> str:
    with tempfile.TemporaryDirectory(prefix="local-report-doc-") as temp_dir:
        temp_path = Path(temp_dir)
        converted = temp_path / f"{path.stem}.docx"
        if _convert_with_textutil(path, converted, notes):
            notes.append("旧版 .doc 已通过 textutil 转换为临时 .docx 后解析。")
            return _parse_word(converted, notes)

        if _convert_with_soffice(path, temp_path, notes):
            candidates = sorted(temp_path.glob("*.docx"))
            if candidates:
                notes.append("旧版 .doc 已通过 LibreOffice 转换为临时 .docx 后解析。")
                return _parse_word(candidates[0], notes)

    notes.append("旧版 .doc 自动转换失败，请确认文件未加密、未损坏，或另存为 .docx 后重试。")
    return ""


def _convert_with_textutil(source: Path, target: Path, notes: list[str]) -> bool:
    executable = shutil.which("textutil")
    if not executable:
        notes.append("未检测到 textutil，跳过 .doc 的 textutil 转换。")
        return False
    return _run_converter(
        [executable, "-convert", "docx", "-output", str(target), str(source)],
        target,
        "textutil",
        notes,
    )


def _convert_with_soffice(source: Path, output_dir: Path, notes: list[str]) -> bool:
    executable = shutil.which("soffice") or shutil.which("libreoffice")
    if not executable:
        notes.append("未检测到 LibreOffice/soffice，跳过 .doc 的 LibreOffice 转换。")
        return False
    expected = output_dir / f"{source.stem}.docx"
    return _run_converter(
        [executable, "--headless", "--convert-to", "docx", "--outdir", str(output_dir), str(source)],
        expected,
        "LibreOffice",
        notes,
    )


def _run_converter(command: list[str], expected_output: Path, name: str, notes: list[str]) -> bool:
    try:
        result = subprocess.run(
            command,
            capture_output=True,
            check=False,
            text=True,
            timeout=60,
        )
    except Exception as exc:
        notes.append(f"{name} 转换 .doc 失败：{exc}")
        return False

    if result.returncode == 0 and expected_output.exists() and expected_output.stat().st_size > 0:
        return True

    detail = (result.stderr or result.stdout or "").strip()
    if detail:
        notes.append(f"{name} 转换 .doc 未成功：{detail[:240]}")
    else:
        notes.append(f"{name} 转换 .doc 未成功。")
    return False


def _parse_word(path: Path, notes: list[str]) -> str:
    document = Document(path)
    parts: list[str] = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            parts.append(text)
    for table_index, table in enumerate(document.tables, start=1):
        parts.append(f"\n[表格 {table_index}]")
        for row in table.rows:
            cells = [_clean_cell(cell.text) for cell in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))
    notes.append("Word 文档解析完成，已读取段落和表格。")
    return "\n".join(parts)


def _parse_pdf(path: Path, notes: list[str]) -> str:
    pymupdf_text = _parse_pdf_with_pymupdf(path)
    if len(pymupdf_text.strip()) >= 500:
        notes.append("PDF 文本层解析完成。")
        return pymupdf_text

    notes.append("PDF 文本层内容较少，已切换 pdfplumber 备选解析。")
    plumber_text = _parse_pdf_with_pdfplumber(path)
    if len(plumber_text.strip()) > len(pymupdf_text.strip()):
        candidate = plumber_text
    else:
        candidate = pymupdf_text

    if len(candidate.strip()) >= 300:
        return candidate

    notes.append("PDF 可能是扫描件，已尝试 OCR。")
    ocr_text = _ocr_pdf(path, notes)
    return ocr_text if len(ocr_text.strip()) > len(candidate.strip()) else candidate


def _parse_pdf_with_pymupdf(path: Path) -> str:
    parts: list[str] = []
    with fitz.open(path) as doc:
        for page_index, page in enumerate(doc, start=1):
            text = page.get_text("text").strip()
            if text:
                parts.append(f"\n[第 {page_index} 页]\n{text}")
    return "\n".join(parts)


def _parse_pdf_with_pdfplumber(path: Path) -> str:
    parts: list[str] = []
    with pdfplumber.open(path) as pdf:
        for page_index, page in enumerate(pdf.pages, start=1):
            text = page.extract_text() or ""
            if text.strip():
                parts.append(f"\n[第 {page_index} 页]\n{text.strip()}")
            tables = page.extract_tables() or []
            for table_index, table in enumerate(tables, start=1):
                parts.append(f"[第 {page_index} 页表格 {table_index}]")
                for row in table:
                    parts.append(" | ".join(_clean_cell(cell or "") for cell in row))
    return "\n".join(parts)


def _ocr_pdf(path: Path, notes: list[str]) -> str:
    if not _tesseract_available():
        notes.append("未检测到 tesseract，扫描 PDF OCR 已跳过。")
        return ""

    parts: list[str] = []
    with fitz.open(path) as doc:
        for page_index, page in enumerate(doc, start=1):
            pix = page.get_pixmap(matrix=fitz.Matrix(2, 2), alpha=False)
            image = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
            text = _image_to_string(image).strip()
            if text:
                parts.append(f"\n[OCR 第 {page_index} 页]\n{text}")
    notes.append("扫描 PDF OCR 完成。")
    return "\n".join(parts)


def _ocr_image(path: Path, notes: list[str]) -> str:
    if not _tesseract_available():
        notes.append("未检测到 tesseract，图片 OCR 无法执行。")
        return ""
    image = Image.open(path)
    text = _image_to_string(image)
    notes.append("图片 OCR 完成。")
    return text


def _tesseract_available() -> bool:
    return shutil.which("tesseract") is not None


def _image_to_string(image: Image.Image) -> str:
    try:
        return pytesseract.image_to_string(image, lang="chi_sim+eng")
    except Exception:
        return pytesseract.image_to_string(image, lang="eng")


def _clean_cell(value: str) -> str:
    return re.sub(r"\s+", " ", value).strip()


def _clean_text(value: str) -> str:
    value = value.replace("\x00", " ")
    value = re.sub(r"[ \t]+", " ", value)
    value = re.sub(r"\n{3,}", "\n\n", value)
    return value.strip()


def detect_language(text: str) -> str:
    sample = text[:4000]
    if not sample:
        return "unknown"
    cjk = len(re.findall(r"[\u4e00-\u9fff]", sample))
    latin = len(re.findall(r"[A-Za-z]", sample))
    if cjk >= max(20, latin * 0.15):
        return "zh"
    if latin >= 80:
        return "en"
    return "unknown"
