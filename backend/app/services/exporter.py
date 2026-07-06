from __future__ import annotations

import re
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

from ..config import get_settings


def export_report_docx(report: dict[str, Any]) -> Path:
    settings = get_settings()
    path = settings.exports_dir / f"report-{report['id']}.docx"
    document = _base_document("可研报告精读结论")
    document.add_paragraph(f"文件：{report['filename']}")
    document.add_paragraph(f"语言：{report.get('language', 'unknown')}")
    document.add_paragraph(f"解析说明：{report.get('parser_notes', '')}")
    _add_analysis(document, report.get("analysis") or {})
    _add_excerpt(document, report.get("extracted_text", ""))
    document.save(path)
    return path


def export_comparison_docx(comparison: dict[str, Any]) -> Path:
    settings = get_settings()
    path = settings.exports_dir / f"comparison-{comparison['id']}.docx"
    document = _base_document("跨报告交叉比对结论")
    result = comparison["result"]
    document.add_paragraph(f"报告数量：{result.get('报告数量', len(comparison.get('report_ids', [])))}")
    for section in ["一致性信号", "矛盾点"]:
        _heading(document, section, level=1)
        for item in result.get(section, []):
            document.add_paragraph(str(item), style="List Bullet")

    _heading(document, "预期差地图", level=1)
    table = document.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    headers = ["主题", "市场共识", "报告判断", "预期差/观察"]
    for idx, header in enumerate(headers):
        table.rows[0].cells[idx].text = header
    for item in result.get("预期差地图", []):
        row = table.add_row().cells
        row[0].text = str(item.get("主题", ""))
        row[1].text = str(item.get("市场共识", ""))
        row[2].text = str(item.get("报告判断", ""))
        row[3].text = str(item.get("预期差", item.get("观察", "")))

    _heading(document, "报告摘要", level=1)
    for summary in result.get("报告摘要", []):
        document.add_paragraph(f"{summary.get('报告', '')}：{summary.get('核心观点', '')}", style="List Bullet")
    document.save(path)
    return path


def _base_document(title: str) -> Document:
    document = Document()
    styles = document.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10.5)
    title_para = document.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run(title)
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(31, 78, 121)
    return document


def _add_analysis(document: Document, analysis: dict[str, Any]) -> None:
    if not analysis:
        document.add_paragraph("尚未生成分析结果。")
        return
    _heading(document, "标准摘要", level=1)
    summary = analysis.get("标准摘要", {})
    for key, value in summary.items():
        _kv(document, key, value)

    _heading(document, "报告修改建议与解决方案", level=1)
    revision = analysis.get("报告修改建议与解决方案", {})
    for key, value in revision.items():
        _kv(document, key, value)

    _heading(document, "最新数据补充与核验清单", level=1)
    latest_data = analysis.get("最新数据补充与核验清单", [])
    _kv(document, "数据清单", latest_data)

    _heading(document, "同类企业对比与业务建议", level=1)
    peer_advice = analysis.get("同类企业对比与业务建议", {})
    for key, value in peer_advice.items():
        _kv(document, key, value)

    _heading(document, "经济测算专项审查", level=1)
    audit = analysis.get("经济测算专项审查", {})
    for key, value in audit.items():
        _kv(document, key, value)

    _heading(document, "精读结论", level=1)
    deep = analysis.get("精读结论", {})
    for key, value in deep.items():
        _kv(document, key, value)

    _heading(document, "输出完整性检查", level=1)
    integrity = analysis.get("输出完整性检查", {})
    for key, value in integrity.items():
        _kv(document, key, value)

    _heading(document, "关键词", level=1)
    keywords = analysis.get("关键词", [])
    document.add_paragraph("、".join(k if isinstance(k, str) else str(k[0]) for k in keywords))


def _add_excerpt(document: Document, text: str) -> None:
    _heading(document, "关键原文摘录", level=1)
    document.add_paragraph(
        f"原文提取总长度约 {len(text)} 字。本节按正文章节摘录，普通段落控制在 800 字以内；涉及关键测算表时保留原始行列结构，完整内容以原报告为准。"
    )
    excerpts = _key_excerpts(text)
    if not excerpts:
        document.add_paragraph(text[:5000] + ("..." if len(text) > 5000 else "") or "无可用原文。")
        return
    for item in excerpts:
        _add_excerpt_item(document, item["label"], item["text"])


def _key_excerpts(text: str) -> list[dict[str, str]]:
    body_start = _body_start(text)
    groups = [
        ("经营模式", ["主营业务模式", "业务主体、模式及业务开展计划"], ["目标计划", "经营效益分析"]),
        ("假设条件", ["投资测算原则和基本假设", "基本假设"], ["业务收益测算", "风险评估"]),
        ("税收与收益测算", ["业务收益测算", "增值税"], ["风险评估", "风险评估及应对措施"]),
        ("所得税与净利润", ["所得税", "净利润"], ["风险评估", "风险评估及应对措施"]),
        ("IRR与收益", ["IRR", "投资收益率", "净利润"], ["风险评估", "风险评估及应对措施"]),
        ("补贴与现金流", ["政府补助", "180天", "资金成本"], ["业务毛利", "风险评估"]),
        ("政策依据", ["政策措施", "商贸发", "成商务"], ["股东介绍", "发展前景评估"]),
    ]
    excerpts: list[dict[str, str]] = []
    seen: set[str] = set()
    for label, terms, end_terms in groups:
        snippet = _chapter_or_keyword_excerpt(text, body_start, terms, end_terms)
        if snippet and snippet not in seen:
            seen.add(snippet)
            excerpts.append({"label": label, "text": snippet})
    for label, terms in [
        ("财务测算表格", ["服务收入", "政府补助", "所得税", "投资收益率"]),
        ("风险评估表格", ["税收政策风险", "国际政策及壁垒"]),
    ]:
        snippet = _table_excerpt_by_terms(text, terms)
        if snippet and snippet not in seen:
            seen.add(snippet)
            excerpts.append({"label": label, "text": snippet})
    return excerpts


def _body_start(text: str) -> int:
    patterns = [
        r"\n公司设立概述\s*\n",
        r"\n背景及目标\s*\n",
        r"\n一、?\s*公司设立概述\s*\n",
    ]
    for pattern in patterns:
        match = re.search(pattern, text)
        if match:
            return match.start()
    return min(len(text), 1200)


def _chapter_or_keyword_excerpt(text: str, body_start: int, terms: list[str], end_terms: list[str]) -> str:
    match = _first_match(text, terms, start=body_start)
    if match is None:
        return ""
    match_start, match_end = match
    start = _line_start(text, match_start)
    end = _next_term_pos(text, end_terms, match_end)
    if end is None or end <= start:
        end = min(len(text), match_end + 1000)
    snippet = text[start:end].strip()
    snippet = _clean_excerpt(snippet)
    return _limit_excerpt(snippet, 800)


def _first_match(text: str, terms: list[str], start: int = 0):
    for term in terms:
        match = re.search(re.escape(term), text[start:], flags=re.I)
        if match:
            absolute_start = start + match.start()
            absolute_end = start + match.end()
            return absolute_start, absolute_end
    return None


def _line_start(text: str, pos: int) -> int:
    previous_newline = text.rfind("\n", 0, pos)
    return 0 if previous_newline < 0 else previous_newline + 1


def _next_term_pos(text: str, terms: list[str], start: int) -> int | None:
    positions = []
    for term in terms:
        match = re.search(re.escape(term), text[start:], flags=re.I)
        if match:
            positions.append(start + match.start())
    return min(positions) if positions else None


def _clean_excerpt(value: str) -> str:
    value = re.sub(r"\[Heading(?:\s+\d+)?\]\s*", "", value)
    value = re.sub(r"\n{3,}", "\n\n", value)
    return value.strip()


def _limit_excerpt(value: str, limit: int) -> str:
    if len(value) <= limit:
        return value
    cut = value[:limit]
    sentence_end = max(cut.rfind("。"), cut.rfind("；"), cut.rfind("\n"))
    if sentence_end > limit * 0.6:
        cut = cut[: sentence_end + 1]
    return cut.rstrip() + "..."


def _table_excerpt_by_terms(text: str, terms: list[str]) -> str:
    for block in _table_blocks(text):
        if all(term.lower() in block.lower() for term in terms):
            return block
    return ""


def _table_blocks(text: str) -> list[str]:
    lines = text.splitlines()
    blocks: list[str] = []
    index = 0
    while index < len(lines):
        line = lines[index].strip()
        if not _is_table_marker(line):
            index += 1
            continue

        block = [line]
        index += 1
        while index < len(lines):
            row = lines[index].strip()
            if not row:
                index += 1
                break
            if _is_table_marker(row):
                break
            if "|" not in row:
                break
            block.append(row)
            index += 1
        if len(block) > 1:
            blocks.append("\n".join(block))
    return blocks


def _add_excerpt_item(document: Document, label: str, text: str) -> None:
    para = document.add_paragraph()
    run = para.add_run(f"{label}：")
    run.bold = True

    segments = _split_excerpt_segments(text)
    if not segments:
        return

    if len(segments) == 1 and segments[0]["type"] == "text":
        para.add_run(segments[0]["value"])
        return

    for segment in segments:
        if segment["type"] == "text":
            value = segment["value"].strip()
            if value:
                document.add_paragraph(value)
        elif segment["type"] == "table":
            marker = segment["marker"]
            marker_para = document.add_paragraph()
            marker_run = marker_para.add_run(marker)
            marker_run.bold = True
            _add_word_table(document, segment["rows"])


def _split_excerpt_segments(text: str) -> list[dict[str, Any]]:
    lines = text.splitlines()
    segments: list[dict[str, Any]] = []
    prose: list[str] = []
    index = 0

    def flush_prose() -> None:
        if prose:
            value = "\n".join(prose).strip()
            if value:
                segments.append({"type": "text", "value": value})
            prose.clear()

    while index < len(lines):
        line = lines[index].strip()
        if not _is_table_marker(line):
            prose.append(lines[index])
            index += 1
            continue

        flush_prose()
        marker = line.strip("[]")
        rows: list[list[str]] = []
        index += 1
        while index < len(lines):
            row = lines[index].strip()
            if not row:
                index += 1
                break
            if _is_table_marker(row):
                break
            if "|" not in row:
                prose.append(lines[index])
                index += 1
                break
            rows.append([_normalize_table_cell(cell) for cell in row.split("|")])
            index += 1
        if rows:
            segments.append({"type": "table", "marker": marker, "rows": rows})
        else:
            prose.append(line)

    flush_prose()
    return segments


def _add_word_table(document: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    column_count = max(len(row) for row in rows)
    table = document.add_table(rows=0, cols=column_count)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    for row_index, row_values in enumerate(rows):
        table_row = table.add_row()
        cells = table_row.cells
        _prevent_row_split(table_row)
        if row_index == 0:
            _repeat_table_header(table_row)
        for cell_index in range(column_count):
            value = row_values[cell_index] if cell_index < len(row_values) else ""
            cell = cells[cell_index]
            cell.text = value
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                if cell_index > 0 and _looks_numeric(value):
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.size = Pt(8 if column_count >= 6 else 9)
                    if row_index == 0:
                        run.bold = True
    document.add_paragraph()


def _repeat_table_header(row: Any) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:tblHeader")) is not None:
        return
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def _prevent_row_split(row: Any) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:cantSplit")) is not None:
        return
    cant_split = OxmlElement("w:cantSplit")
    cant_split.set(qn("w:val"), "true")
    tr_pr.append(cant_split)


def _is_table_marker(value: str) -> bool:
    return bool(re.match(r"^\[(?:表格\s*\d+|第\s*\d+\s*页表格\s*\d+)\]$", value))


def _normalize_table_cell(value: str) -> str:
    return re.sub(r"\s+", " ", value).strip()


def _looks_numeric(value: str) -> bool:
    return bool(re.search(r"^[\d,.%（）()半年年\s-]*$", value.strip()))


def _heading(document: Document, text: str, level: int) -> None:
    document.add_heading(text, level=level)


def _kv(document: Document, key: str, value: Any) -> None:
    para = document.add_paragraph()
    run = para.add_run(f"{key}：")
    run.bold = True
    if isinstance(value, list):
        if not value:
            para.add_run("无")
        else:
            for item in value:
                if isinstance(item, dict):
                    _dict_list_item(document, item)
                else:
                    document.add_paragraph(str(item), style="List Bullet")
    elif isinstance(value, dict):
        para.add_run("")
        for nested_key, nested_value in value.items():
            _kv(document, nested_key, nested_value)
    else:
        para.add_run(str(value))


def _dict_list_item(document: Document, item: dict[str, Any]) -> None:
    primary_key = next(iter(item.keys()), "")
    primary_value = item.get(primary_key, "")
    para = document.add_paragraph(style="List Bullet")
    if primary_key:
        run = para.add_run(f"{primary_key}：")
        run.bold = True
        para.add_run(str(primary_value))
    for key, value in list(item.items())[1:]:
        child = document.add_paragraph()
        child.paragraph_format.left_indent = Pt(18)
        run = child.add_run(f"{key}：")
        run.bold = True
        child.add_run(_format_value(value))


def _format_value(value: Any) -> str:
    if isinstance(value, list):
        return "；".join(_format_value(item) for item in value)
    if isinstance(value, dict):
        return "；".join(f"{key}：{_format_value(item)}" for key, item in value.items())
    return str(value)
